# extracting/msft_extractor.py

import tempfile
from pathlib import Path
from typing import List

import pythoncom
import win32com.client

from docx import Document
import mammoth
from striprtf.striprtf import rtf_to_text

from .extractors import FileTextExtractor

class WordFileTextExtractor(FileTextExtractor):
    """
    Windows-friendly text extractor for Word formats.
    - DOCX/DOCM: mammoth -> markdown (fallback python-docx)
    - DOC/RTF:   convert via Word COM to TXT (fast, reliable) then read
                 (fallback to pandoc or striprtf if Word isn't installed)
    """
    file_extensions: List[str] = ["docx", "docm", "doc", "rtf"]

    def __init__(self, use_mammoth: bool = True, use_word_com: bool = True,
                 pandoc_path: str | None = None):
        super().__init__()
        self.use_mammoth  = use_mammoth
        self.use_word_com = use_word_com
        self.pandoc_path  = pandoc_path

    def __call__(self, path: str) -> str:
        ext = Path(path).suffix.lower().lstrip(".")
        if ext in ("docx", "docm"):
            return self._extract_docx(path)
        elif ext in ("doc", "rtf"):
            return self._extract_legacy(path, ext)
        else:
            raise ValueError(f"Unsupported Word extension: {ext}")

    # ---------- helpers ----------
    def _extract_docx(self, path: str) -> str:
        if self.use_mammoth:
            try:
                with open(path, "rb") as f:
                    return mammoth.convert_to_markdown(f).value
            except Exception:
                pass  # fall through to python-docx

        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)

    def _extract_legacy(self, path: str, ext: str) -> str:
        if self.use_word_com:
            try:
                return self._word_com_to_txt(path)
            except Exception:
                pass  # fall back

        # Fallbacks
        if ext == "rtf":
            with open(path, "r", encoding="latin-1", errors="ignore") as f:
                return rtf_to_text(f.read())

        if self.pandoc_path:
            return self._pandoc_to_txt(path)

        raise RuntimeError("No viable method to extract text from legacy Word file on Windows.")

    def _word_com_to_txt(self, path: str) -> str:
        """
        Use Microsoft Word via COM to SaveAs TXT, then read.
        """
        # Constants from Word Object Model (avoid importing win32com.constants each call)
        wdFormatText = 2
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        txt = ""
        try:
            doc = word.Documents.Open(Path(path).absolute().__str__(), ReadOnly=True)
            with tempfile.TemporaryDirectory() as td:
                out_txt = Path(td) / (Path(path).stem + ".txt")
                doc.SaveAs2(str(out_txt), FileFormat=wdFormatText, Encoding=65001)  # UTF-8
                doc.Close()
                with open(out_txt, "r", encoding="utf-8", errors="ignore") as f:
                    txt = f.read()
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
        return txt

    def _pandoc_to_txt(self, path: str) -> str:
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tmp:
            out = tmp.name
        cmd = [self.pandoc_path, path, "-t", "plain", "-o", out]
        subprocess.run(cmd, check=True)
        with open(out, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
